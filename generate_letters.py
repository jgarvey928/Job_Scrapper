import csv
import google.generativeai as genai
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import time

GOOGLE_API_KEY = "AIzaSyAtLFVTbGsSWx30-BMUhBHP-SesZGMF55c"

def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def write_docx(file_path, content):
    doc = docx.Document()
    paragraph = doc.add_paragraph(content)
    
    # Set font to Arial and size to 12
    run = paragraph.runs[0]
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Add spacing between paragraphs
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    
    doc.save("letters/"+file_path)

def analyze_text_from_file(system_instructions, job_data, cl_template):

    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-pro-latest')

    try:
        # Combine instructions and content into a single string
        prompt = system_instructions + "\n\n" + job_data + "\n\n" + cl_template  # Add a separator for clarity

        response = model.generate_content(prompt)
        print("\n$$$ Response Complete $$$\n")
        return response.text
    except Exception as e:
        print(f"Error during Gemini API call: {e}")
        return None

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

# Function to check if a string contains any of the keywords
def contains_keywords(text):
    
    # List of keywords to identify the requirements section
    keywords = [
        "entry level", "entry-level", "early career", 
        "early-career", "junior", "jr", "1",
    ]   
    
    # Check if i... ( the char i followed by any numbers or letters) is in the text
    if re.search(r'\bi\b', text.lower()):
        return True
    
    for keyword in keywords:
        if keyword.lower() in text.lower() and not "senior" in text.lower():
            return True
    return False 

# Custom key function to prioritize job titles containing "jr"
def sort_key(job):
    #TODO: Fix this return sorting logic
    return( 0 if("entry level" in job["Job Title"].lower() or
                 "entry-level" in job["Job Title"].lower() or
                 "early career" in job["Job Title"].lower() or
                 "early-career" in job["Job Title"].lower() or
                 "junior" in job["Job Title"].lower() or
                 "jr" in job["Job Title"].lower() )
            else 1)

def generate_letters(file_name, letter_count):
    
    input_file = 'scrapped_data/'+file_name+'.csv'
    job_descriptions = []
    top_jobs = []

    # Read the CSV file
    with open(input_file, mode='r', encoding='utf-8') as infile:
        reader = csv.DictReader(infile)
        count = 0

        for row in reader:
            # Extract the required fields
            job_info = {
                "Job Title": row["Job Title"],
                "Company": row["Company"],
                "Cleaned": row["Cleaned"]
            }
            # print(str(job_info)+"\n")
            job_descriptions.append(job_info)

            if(contains_keywords(job_info["Job Title"]) and count < letter_count ):
                top_jobs.append(job_info)
                count += 1

    # Example usage:
    cl_file_path = "letters/cover_letter_template.txt"

    system_instructions = """I am a recent software engineering graduate looking for work. You are a career coach 
                            with over 20 years experience whose helping me fill in my cover letter for a job.
                            Theres a section in square brackets [] in the cover letter that I want you to fill in. 
                            Fill in the section as instructed inside the brackets and based on the job description. 
                            Make it sound human but professional and enthusiastic. Only fill in the bracketed sections
                            leave the rest alone. Heres the job description followed by my cover letter I want you to help complete..."""
    
    try:
        cl_text_content = read_text_file(cl_file_path)
    except FileNotFoundError as e:
        print(f"Error: File not found at {e.filename}")
        return None
    except Exception as e:
        print(f"Error reading file: {e}")
        return None

    # Sort the list using the custom key function
    sorted_top_jobs = sorted(top_jobs, key=sort_key)
    for data in sorted_top_jobs:
        job_data_content = data["Job Title"] + "\n" + data["Company"] + "\n" + data["Cleaned"]
        print("Entry Level Job:  " + data["Job Title"] + "   " + data["Company"])
        response = analyze_text_from_file(system_instructions, job_data_content, cl_text_content)
        time.sleep(2)
        if response:
            # Replace spaces and special characters with no spaces in the file name
            job_title = re.sub(r'\W+', '', data["Job Title"].replace(" ", ""))
            company = re.sub(r'\W+', '', data["Company"].replace(" ", ""))
            file_name = f"{job_title}_{company}_letter.docx"
            write_docx(file_name, response)
        else:
            print("Failed to get a response from Gemini.")


if __name__ == '__main__':
    file = 'se_phx_entry_onsite_month_02_23_25'
    generate_letters(file, 10)